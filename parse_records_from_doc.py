from docx import Document
import pickle
import re

document = Document('rigaspool.docx')

data = {year: {} for year in range(2005, 2007)}


def get_week(num):
    if num < 6:
        wk = 23 - num
    else:
        wk = 24 - num
    return wk

def get_week2(num):
    if num < 9:
        wk = 26 - num
    elif num < 29:
        wk = 29 - num
    elif num < 43:
        wk = 58 - num
    elif num < 63:
        wk = 62 - num
    else:
        wk = 0
    return wk



# for idx_t, table in enumerate(document.tables[1:]):
#     year = 2018 - idx_t
#     player = {}
#     for idx_r, row in enumerate(table.rows):
#         for idx_c, cell in enumerate(row.cells):
#             if cell.text:
#                 if idx_r == 0 and idx_c != 0:
#                     player[idx_c] = cell.text
#                     data[year][cell.text] = {}
#                 elif idx_r == 1 or idx_r == 6:
#                     pass
#                 else:
#                     if idx_c != 0:
#                         record = [re.sub("[^0-9]", "",s) for s in re.split('-|=',cell.text)]
#                         int_record = [int(total) if total else 0 for total in record]
#                         data[year][player[idx_c]][get_week(idx_r)] = int_record

# with open('nfl_picks_history.pl','wb') as f:
#     pickle.dump(data, f)
need_name = True
year = 2006
for idx, paragraph in enumerate(document.paragraphs[-62:]):
    if paragraph.text:
        line = re.split('\s', paragraph.text)
        line = [l for l in line if l][-6:]
        if line and line[0] == 'MJR' and need_name:
            need_name = False
            player = line
            data[2005] = {p:{} for p in player}
            data[2006] = {p:{} for p in player}
        if line == ['2005']:
            year = 2005
        if 'MJR' not in line and '2005' not in line and idx > 4 and idx != 9 and idx != 35:
            for p_idx, rec in enumerate(line):
                if len(rec) < 6:
                    record = [re.sub("[^0-9]", "",s) for s in re.split('-|=',rec)]
                    int_record = [int(total) if total else 0 for total in record]
                    data[year][player[p_idx]][get_week2(idx)] = int_record


data[2005]['TJR'][4] = [10, 4]

with open('nfl_picks_history.pl','wb') as f:
     pickle.dump(data, f)

        